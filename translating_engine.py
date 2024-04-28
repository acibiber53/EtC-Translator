from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import selenium.common.exceptions as sce
from time import sleep, time
from datetime import date, timedelta
import re
from docx import Document
import os
import subprocess
from webdriver_manager.chrome import ChromeDriverManager
from news import News
import logging
# from fake_useragent import UserAgent
from news_outlets_xpaths import headers, bodies, outlet_chinese_names
from bs4 import BeautifulSoup
from trello_controller import (
    create_card_with_trello,
    TrelloController,
    upload_daily_news_to_trello,
)
from gdapi_controller import GoogleDriveAPIController as GDAPI

logging.basicConfig(level=logging.DEBUG)


class Translator:
    news_uploaded = list()

    def __init__(self, tengine="sogou"):
        # ua = UserAgent()
        # userAgent = ua.random
        userAgent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
        self.driver = self.open_browser(userAgent)
        self.preferred_translation_engine = tengine
        self.translator = self.open_browser(userAgent)
        self.stinput = None
        self.stoutput = None
        self.translator_settings()
        sleep(2)
        self.output_prefix = self.prepare_prefix()
        self.output_directory = self.output_prefix[:-3] + "\\"
        self.current_news = None

    def translator_settings(self):
        if self.preferred_translation_engine == "sogou":
            self.translator.get("https://fanyi.sogou.com/")
            self.stoutput = self.translator.find_element_by_class_name("output")
            self.stinput = self.translator.find_element_by_id("trans-input")
        elif self.preferred_translation_engine == "baidu":
            self.translator.get("https://fanyi.baidu.com/#en/zh/")
            try:
                popup_element = self.translator.find_element_by_class_name(
                    "desktop-guide-close"
                )
                popup_element.click()
            except:
                print("NO popups I guess")
            try:
                self.stinput = self.translator.find_element_by_id(
                    "baidu_translate_input"
                )
                self.stoutput = self.translator.find_element_by_class_name(
                    "trans-right"
                )
            except sce.NoSuchElementException as e:
                print(
                    "Translation Engine might have changed the element structure. Update needed!"
                )

    @staticmethod
    def prepare_prefix():
        """
            This method prepares the date prefix for the document names. It is in the format of "XXXX.XX.XX - ". After
            this comes the header of the news.
        :return:
        """
        return re.sub("-", ".", str(date.today())) + " - "

    @staticmethod
    def open_browser(userA):
        """
            This method opens a webdriver in a preset options environment. Now it only opens the driver maximized and
            ignores the ssl errors.

            Some other possible arguments also left in the comments for future reference. It is possible to use it in
            headless mode with the arguments in the comment.

        :return:  webdriver
        """
        """
        # options for headless Chrome instance that passes automation detection
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--window-size=1280x800')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('--ignore-certificate-errors')
        options.add_argument('--ignore-ssl-errors')

        driver = webdriver.Chrome(options=options)  
        
        """
        options = webdriver.ChromeOptions()
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("--ignore-ssl-errors")
        options.add_argument("start-maximized")
        options.add_argument(
            f"user-agent={userA}"
        )  # https://stackoverflow.com/questions/49565042/way-to-change-google-chrome-user-agent-in-selenium/49565254#49565254
        options.add_argument(
            "--disable-blink-features=AutomationControlled"
        )  # https://stackoverflow.com/questions/53039551/selenium-webdriver-modifying-navigator-webdriver-flag-to-prevent-selenium-detec/53040904#53040904

        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option("useAutomationExtension", False)

        # Added version to ChromeDriverManager, as the changes to version url as of 2023.07.19 made it stop working.
        # Also, for every other library we are using old versions, why were we using new ChromeDriver each time?
        # Possible Answer: Because of websites' user agent detection?
        try:
            driver = webdriver.Chrome(ChromeDriverManager().install(), options=options)
        except sce.SessionNotCreatedException as error:
            print(error)
            os.system("pause")
        sleep(1)

        driver.execute_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        # driver.execute_cdp_cmd('Network.setUserAgentOverride', {"userAgent": userA})
        print(driver.execute_script("return navigator.userAgent;"))
        sleep(1)

        return driver

    def close_driver(self):
        """
        This method closes all the drivers that are opened by the class.
        """
        self.driver.quit()
        self.translator.quit()

    def find_news_outlet(self, source_link):
        tmp = source_link.split("/")[2].split(".")
        if tmp[0] == "www":
            news_outlet = tmp[1]
        else:
            news_outlet = tmp[0]
        return news_outlet

    def parse_link(self):
        is_not_found = 0
        try:
            header = self.driver.find_element_by_xpath(
                headers.get(self.current_news.news_outlet, "//h1")
            ).text
        except sce.NoSuchElementException as error:
            print(
                f"We got an error message when searching for header:\n{error}\nTo be able to continue our "
                f"work, we are selecting the header from the most common xpath, //h1."
            )
            header = self.driver.find_element_by_xpath("//h1").text
            is_not_found = 1

        try:
            body = self.driver.find_elements_by_xpath(
                bodies.get(self.current_news.news_outlet, "//p")
            )
        except sce.NoSuchElementException as error:
            print(
                f"We got an error message when searching for body:\n{error}\nTo be able to continue our "
                f"work, we are selecting the body from the most common xpath, //p."
            )
            body = self.driver.find_elements_by_xpath("//p")
            is_not_found = 1

        if is_not_found:
            print("We stopped because we couldn't find header or body!")
            os.system("pause")

        self.current_news.title_english = re.sub(r'[\\/:"*?<>|]+', "-", header)

        self.current_news.body_english = "\n".join(
            [item.text for item in body if item.text.strip()]
        )

        self.current_news.full_text_english = "\n".join(
            [self.current_news.title_english, self.current_news.body_english]
        )
        self.current_news.length_english = len(self.current_news.full_text_english)

    def translate_with_sogou(self):
        fulltext = self.current_news.full_text_english

        paragraphs = fulltext.split("\n")
        par_count = len(paragraphs)
        par_point = 0
        output = ""

        while par_point < par_count:
            translate_text = ""
            while par_point < par_count:
                if len(translate_text) + len(paragraphs[par_point]) > 5000:
                    break
                if translate_text:
                    translate_text = "\n".join([translate_text, paragraphs[par_point]])
                else:
                    translate_text = paragraphs[par_point]
                par_point += 1
            # Copying body to clipboard
            subprocess.run(
                ["clip.exe"], input=translate_text.strip().encode("utf-16"), check=True
            )
            # Pasting it into the translation input
            self.stinput.send_keys(Keys.CONTROL, "v")
            sleep(2)
            if output:
                output = "\n".join([output, self.stoutput.text])
            else:
                output = self.stoutput.text

            # Cleaning the text area for next translation
            clean_button = self.translator.find_element_by_xpath(
                "//div[@class='trans-con']/span"
            )
            clean_button.click()

            return output

    def translate_with_baidu(self):
        fulltext = self.current_news.full_text_english

        paragraphs = fulltext.split("\n")
        par_count = len(paragraphs)
        par_point = 0
        output = ""

        while par_point < par_count:
            translate_text = ""
            while par_point < par_count:
                if len(translate_text) + len(paragraphs[par_point]) > 5000:
                    break
                if translate_text:
                    translate_text = "\n".join([translate_text, paragraphs[par_point]])
                else:
                    translate_text = paragraphs[par_point]
                par_point += 1
            # Copying body to clipboard
            subprocess.run(
                ["clip.exe"], input=translate_text.strip().encode("utf-16"), check=True
            )
            # Pasting it into the translation input
            self.stinput.send_keys(Keys.CONTROL, "v")

            translated_pars = False
            while not translated_pars:
                translated_pars = self.stoutput.find_elements_by_class_name(
                    "ordinary-output.target-output.clearfix"
                )
                sleep(1)

            tmp_res = "\n".join([elem.text for elem in translated_pars])

            if output:
                output = "\n".join([output, tmp_res])
            else:
                output = tmp_res

            # Cleaning the text area for next translation
            clean_button = self.translator.find_element_by_class_name(
                "textarea-clear-btn"
            )
            clean_button.click()

        return output

    def translation_engine_selector(self):
        if self.preferred_translation_engine == "sogou":
            return self.translate_with_sogou()
        elif self.preferred_translation_engine == "baidu":
            return self.translate_with_baidu()

    def fix_non_chinese_chars(self, text):
        """
        This method fixes non-chinese punctuation marks right after the translation finishes.
        :param text:
        :return:
        """
        text = text.replace("(", "（")
        text = text.replace(")", "）")
        text = text.replace(":", "：")
        text = text.replace('"', "”")
        text = text.replace("'", "’")
        text = re.sub(
            r"/(?<!\d)[.]/g", "。", text
        )  # This line checks if there is a digit before the dot, then changes it.
        text = text.replace(",", "，")
        return text

    def fix_translation(self, text):
        """
        This method gets in the translated text and does some last minute changes to it.
        :param text:
        :return:
        """
        text = text.replace("图尔基耶", "土耳其")
        text = text.replace("Turkiye", "土耳其")
        text = text.replace("伊斯坦堡", "伊斯坦布尔")
        text = text.replace("卫生防护中心", "共和人民党")
        text = text.replace("Imamoglu", "伊马毛卢")
        text = text.replace("Kilicdaroglu", "克勒赤达罗卢")
        text = text.replace("基利奇达罗格鲁", "克勒赤达罗卢")
        text = text.replace("辛塞克", "希姆谢克")
        text = text.replace("IYI党", "好党")
        return text

    def add_chinese_news_outlet(self):
        """
        This method adds the line below to the very beginning of the text translation. Its format can be changed later
        on.
        据《某个通讯社》几月几日消息，
        :return:
        """
        if outlet_chinese_names.get(self.current_news.news_outlet):
            outlet_to_add = outlet_chinese_names.get(self.current_news.news_outlet)
        else:
            outlet_to_add = "UNKNOWN NEWS OUTLET"
        yesterday = date.today() - timedelta(days=1)
        text_to_add = (
            "\n据"
            + outlet_to_add
            + str(yesterday.month)
            + "月"
            + str(yesterday.day)
            + "日消息，"
        )
        self.current_news.body_chinese = text_to_add + self.current_news.body_chinese

    def translate(self):
        """
            Main translation method. It receives full text from current_news variable, then pastes it to
            the main translation webelement.

            Translation engine automatically translates it, then it collects the results from the output webelement.

            At the end it clears the input area for the next translation.
        :return:
        """

        output = self.translation_engine_selector()
        output = self.fix_non_chinese_chars(output)
        output = self.fix_translation(output)

        self.current_news.length_chinese = len(output)
        logging.debug(f"Original output is like this:\n{output}")
        # Splitting output into smaller pieces
        all_translation = output.split("\n")
        self.current_news.title_chinese = all_translation[0]
        self.current_news.body_chinese = "\n".join(all_translation[1:])
        self.add_chinese_news_outlet()
        logging.debug(f"Translated chinese title is: {self.current_news.title_chinese}")
        logging.debug(f"Translated chinese body is: {self.current_news.body_chinese}")

    def output_news(self, ch_heading=None, ch_body=None):
        if ch_heading is None:
            ch_heading = self.current_news.title_chinese
        if ch_body is None:
            ch_body = self.current_news.body_chinese.split("\n")

        outputfile = Document()

        # Adding content to document
        logging.debug(f"This is Chinese heading for the document: {ch_heading}")
        outputfile.add_heading(ch_heading)
        logging.debug(f"This is Chinese body for the document: {ch_body}")
        for par in ch_body:
            par += "\n"
            outputfile.add_paragraph(par)

        outputfile.add_paragraph(self.driver.current_url + "\n")

        # Making file name
        self.current_news.document_name = self.output_prefix + ch_heading + ".docx"

        # Check if the named folder exists, if not make one
        if not os.path.exists(self.output_directory):
            os.mkdir(self.output_directory)

        self.current_news.document_path = (
            self.output_directory + self.current_news.document_name
        )

        if os.path.exists(self.current_news.document_path):
            os.remove(self.current_news.document_path)

        outputfile.save(self.current_news.document_path)

    def popup_check(self):
        if self.current_news.news_outlet == "apnews":
            try:
                pop_close = self.driver.find_element_by_xpath(
                    "//button[@class='sailthru-overlay-close']"
                )
                pop_close.click()
            except sce.NoSuchElementException as error:
                print(error)
                print("It is not asking for email so we just skip that part")
        elif self.current_news.news_outlet == "ahvalnews":
            try:
                pop_close = self.driver.find_element_by_xpath("//a[@class='close']")
                webdriver.ActionChains(self.driver).send_keys(Keys.ESCAPE).perform()
            except sce.NoSuchElementException as error:
                print(error)

    def fix_turkish_characters(self):
        text = self.current_news.full_text_english
        text = text.replace("İ", "I")
        text = text.replace("ı", "i")
        text = text.replace("Ç", "C")
        text = text.replace("ç", "c")
        text = text.replace("Ğ", "G")
        text = text.replace("ğ", "g")
        text = text.replace("Ö", "O")
        text = text.replace("ö", "o")
        text = text.replace("Ş", "S")
        text = text.replace("ş", "s")
        text = text.replace("Ü", "U")
        text = text.replace("ü", "u")
        self.current_news.full_text_english = text

    def translate_main(self, link):
        self.current_news = News(link)
        self.driver.get(self.current_news.source_link)
        i = 0
        while self.driver.execute_script("return document.readyState;") != "complete":
            sleep(1)
            i += 1
            print("Checking page readiness")
            if i == 150:
                self.driver.execute_script("window.stop();")
                break
        self.current_news.news_outlet = self.find_news_outlet(
            self.current_news.source_link
        )
        self.popup_check()
        self.parse_link()
        self.fix_turkish_characters()
        self.translate()
        self.output_news()
        self.news_uploaded.append(self.current_news)

    def display_news_uploaded(self):
        for news in self.news_uploaded:
            print(
                news.title_english,
                news.title_chinese,
                news.body_chinese,
                news.source_link,
                sep="\n",
            )
            print()


def htm_to_urllist(doc_name):
    """This method opens a predetermined htm file that is extracted from the bookmarks of chrome, and
    lists every link inside of it.

    Preparation for the htm file should be done well. Steps for that:
    1 - Put all your websites you want to translate into one bookmark folder. Ctrl + Shift + D for short.
    2 - Open the file, find your folder, copy the content.
    3 - Open MS Word, paste the content, save the folder as .htm/ .html"""
    try:
        with open(doc_name, encoding="utf-8") as file:
            soup = BeautifulSoup(file, "lxml")
    except FileNotFoundError:
        print(
            f"We couldn't find your document, please make sure to have a document named {doc_name}. "
            f"Add the required htm file and try again!"
        )
        os.system("pause")
        return -1, -1

    urllist = list()
    url_title_list = list()
    for link in soup.find_all("a"):
        urllist.append(link.get("href"))
        url_title_list.append(link.get_text().replace("\n", " "))
    return url_title_list, urllist


def translate_news(t_engine, url_list, printing_func, window):
    if url_list == -1:
        return

    trello_daily_card = list()
    gdapi = GDAPI()
    trello_instance = TrelloController()
    trs = Translator(t_engine)

    try:
        for index, link in enumerate(url_list):
            # Initial Translation Part
            start_time = time()
            printing_func(f"Translation begins for {link}")
            trs.translate_main(link)
            printing_func(f"Translation ends, it took {time() - start_time} seconds.\n")

            # Creating Trello card template for later use
            trello_daily_card.append(
                (
                    trs.current_news.title_english,
                    trs.current_news.source_link,
                )
            )

            # Uploading to Google
            printing_func("Uploading to Google Drive!")
            (
                news_title,
                trs.current_news.google_upload_link,
            ) = gdapi.docx_to_gdocs_uploader(
                trs.current_news.document_name,
                trs.current_news.document_path,
            )
            printing_func(
                f"Uploaded!\nFile Name: {news_title}\nFile URL: {trs.current_news.google_upload_link}\n"
            )

            # Creating Trello Card
            create_card_with_trello(trs.current_news, printing_func, trello_instance)
            window["-PROG-"].update(index + 1)

    except Exception as error:
        # print(error)
        printing_func(error)

    finally:
        upload_daily_news_to_trello(
            trello_daily_card,
            trs.current_news.translation_date,
            printing_func,
            trello_instance,
        )
        trs.close_driver()


if __name__ == "__main__":
    trans = Translator("sogou")
    url = "https://www.reuters.com/business/energy/turkey-talks-with-exxonmobil-over-multibillion-dollar-lng-deal-ft-reports-2024-04-28/"

    try:
        trans.translate_main(url)
    finally:
        trans.close_driver()
        trans.display_news_uploaded()
